#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
PDF/DOCX extractor for stakeholder discovery.
- Extracts plain text, emails, phones, and links.
- Emits JSON Lines to stdout or a file.

Usage:
  python scripts/pdf_extract.py --paths file1.pdf file2.docx --output out.jsonl
  python scripts/pdf_extract.py --dir ./downloads --recursive
"""

from __future__ import annotations

import argparse
import json
import os
import re
from datetime import datetime, timezone
from typing import List, Dict, Any, Iterable

try:
    from pdfminer.high_level import extract_text as pdf_extract_text
except Exception:
    pdf_extract_text = None  # type: ignore

try:
    import docx  # python-docx
except Exception:
    docx = None  # type: ignore

LINK_PATTERN = re.compile(r"https?://\S+")
EMAIL_PATTERN = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_PATTERN = re.compile(r"(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{2,4}\)?[\s.-]?)?\d{3,4}[\s.-]?\d{3,4}")


def now_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def normalize_space(s: str) -> str:
    return re.sub(r"\s+", " ", s or "").strip()


def extract_from_pdf(path: str) -> str:
    if pdf_extract_text is None:
        raise RuntimeError("pdfminer.six is required. Install with: pip install pdfminer.six")
    return pdf_extract_text(path) or ""


def extract_from_docx(path: str) -> str:
    if docx is None:
        raise RuntimeError("python-docx is required. Install with: pip install python-docx")
    d = docx.Document(path)
    parts: List[str] = []
    for p in d.paragraphs:
        parts.append(p.text)
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def extract_emails(text: str) -> List[str]:
    return sorted(set(EMAIL_PATTERN.findall(text)))


def extract_phones(text: str) -> List[str]:
    raw = PHONE_PATTERN.findall(text)
    cleaned: List[str] = []
    for m in raw:
        digits = re.sub(r"\D", "", m)
        if len(digits) >= 8:
            cleaned.append(normalize_space(m))
    return sorted(set(cleaned))


def extract_links(text: str) -> List[str]:
    # basic link harvest (text-based, not PDF annotations)
    return sorted(set(LINK_PATTERN.findall(text)))


def scan_files(paths: Iterable[str]) -> Iterable[Dict[str, Any]]:
    for path in paths:
        ext = os.path.splitext(path)[1].lower()
        if not os.path.isfile(path):
            continue
        try:
            if ext == ".pdf":
                text = extract_from_pdf(path)
            elif ext == ".docx":
                text = extract_from_docx(path)
            else:
                continue
            text_norm = normalize_space(text)
            item = {
                "path": path,
                "type": ext.lstrip("."),
                "text_size": len(text_norm),
                "emails": extract_emails(text_norm),
                "phones": extract_phones(text_norm),
                "links": extract_links(text_norm),
                "extracted_at": now_iso(),
            }
            yield item
        except Exception as e:
            yield {
                "path": path,
                "type": ext.lstrip("."),
                "error": str(e),
                "extracted_at": now_iso(),
            }


def iter_dir(root: str, recursive: bool) -> List[str]:
    out: List[str] = []
    for base, dirs, files in os.walk(root):
        for f in files:
            if f.lower().endswith((".pdf", ".docx")):
                out.append(os.path.join(base, f))
        if not recursive:
            break
    return out


def build_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(description="Extract text/emails/phones/links from PDF/DOCX")
    p.add_argument("--paths", nargs="*", help="Explicit file paths")
    p.add_argument("--dir", type=str, help="Directory to scan")
    p.add_argument("--recursive", action="store_true", help="Recurse into subdirectories")
    p.add_argument("--output", type=str, help="Write JSONL to this file instead of stdout")
    return p


def main() -> None:
    args = build_parser().parse_args()
    files: List[str] = []
    if args.paths:
        files.extend(args.paths)
    if args.dir:
        files.extend(iter_dir(args.dir, args.recursive))
    if not files:
        print(json.dumps({"error": "No input files provided. Use --paths or --dir"}))
        return

    sink = open(args.output, "w", encoding="utf-8") if args.output else None
    try:
        for rec in scan_files(files):
            line = json.dumps(rec, ensure_ascii=False)
            if sink:
                sink.write(line + "\n")
            else:
                print(line)
    finally:
        if sink:
            sink.close()


if __name__ == "__main__":
    main()